import os
import html
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from django.utils.html import strip_tags

from matrix_fate.matrix_fate_app.models import(
    SahasraraO7, SahasraraP7, SahasraraQ7,
    AdjnaO6, AdjnaP6, AdjnaQ6, VishudkhaO5, VishudkhaP5, VishudkhaQ5,
    AnakhataO4, AnakhataP4, AnakhataQ4, ManipuraO3, ManipuraP3, ManipuraQ3,
    SvadkhistanaO2, SvadkhistanaP2, SvadkhistanaQ2, MuladkharaO1, MuladkharaP1, MuladkharaQ1, TotalO, TotalP, TotalQ, 
)

from matrix_fate.child_app.models import (
    ChildCategory,

    ChildBusinessCard,

    QualitiesRevealedAgeOf20,
    ThirdTalentRevealedAge40,

    ChildOpportunity,

    ChildPointOfComfort,

    MainTaskSoul,
    SoulPastExperiencesWithPeople,
    LessonsFromPastLife,

    ChildDestinyArcana1,
    ChildDestinyArcana2,
    ChildDestinyArcana3,
    
    WhatChildShouldTeachParents,
    WhatMistakesRelationshipParentsChildren,
    WhatShouldComeQualitiesChild,
    MatrixChildProgram,
)

MARKER_MODEL_MAP = {
    "a": [ChildBusinessCard, WhatChildShouldTeachParents],
    "b": [QualitiesRevealedAgeOf20],
    "c": [ThirdTalentRevealedAge40],
    "d": [MainTaskSoul],
    "e": [ChildPointOfComfort],
    "e1": [],
    "e2": [],
    "f": [],
    "g": [],
    "h": [],
    "i": [],
    "a1": [WhatShouldComeQualitiesChild],
    "a2": [ChildOpportunity, WhatMistakesRelationshipParentsChildren],
    "n": [],
    "b1": [],
    "b2": [],
    "m": [],
    "c1": [],
    "c2": [],
    "d1": [LessonsFromPastLife],
    "d2": [SoulPastExperiencesWithPeople],
    "f1": [],
    "f2": [],
    "g1": [],
    "g2": [],
    "h1": [],
    "h2": [],
    "i1": [],
    "i2": [],
    "j": [],
    "k": [],
    "l": [],
    "r": [ChildDestinyArcana1],
    "s": [ChildDestinyArcana2],
    "y": [ChildDestinyArcana3],
    "t": [],
    "u": [],
    "v": [],
    "w": [],
    "x": [],
    "o": [TotalO],
    "o1": [MuladkharaO1],
    "o2": [SvadkhistanaO2],
    "o3": [ManipuraO3],
    "o4": [AnakhataO4],
    "o5": [VishudkhaO5],
    "o6": [AdjnaO6],
    "o7": [SahasraraO7],
    "p": [TotalP],
    "p1": [MuladkharaP1],
    "p2": [SvadkhistanaP2],
    "p3": [ManipuraP3],
    "p4": [AnakhataP4],
    "p5": [VishudkhaP5],
    "p6": [AdjnaP6],
    "p7": [SahasraraP7],
    "q": [TotalQ],
    "q1": [MuladkharaQ1],
    "q2": [SvadkhistanaQ2],
    "q3": [ManipuraQ3],
    "q4": [AnakhataQ4],
    "q5": [VishudkhaQ5],
    "q6": [AdjnaQ6],
    "q7": [SahasraraQ7],
  }


def autofit_table(table):
    try:
        for row in table.rows:
            for cell in row.cells:
                if any(cell._element.xpath('.//w:drawing')):
                    continue
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:tcW {} w:w="5000" w:type="auto"/>'.format(
                        nsdecls('w')))
                )
    except Exception as e:
        print(f"Ошибка в autofit_table: {e}")


def replace_text_in_table(table, replacements):
    autofit_table(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not any(run._element.xpath('.//w:drawing') for run in paragraph.runs):
                    replace_text_with_style(paragraph, replacements)


def replace_text_in_headers(doc, replacements):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if not any(run._element.xpath('.//w:drawing') for run in paragraph.runs):
                replace_text_with_style(paragraph, replacements)


def replace_text_with_style(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)

    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"

        if placeholder in full_text and key.startswith("IMAGE_"):
            image_path = replacements.get(key)
            if image_path and os.path.exists(image_path):
                for run in paragraph.runs:
                    run.text = ""

                try:
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(2.5))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    print(f"Ошибка при вставке изображения: {e}")
            return

        if placeholder in full_text:
            formatted_value = str(value).replace("\n", "\n") if value else ""
            full_text = full_text.replace(placeholder, formatted_value)

            for run in paragraph.runs:
                run.text = ""

            lines = full_text.split("\n")
            for i, line in enumerate(lines):
                run = paragraph.add_run(line)

                if key.startswith("CATEGORY_TITLE"):
                    run.font.size = Pt(22)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(68, 0, 86)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                elif key.startswith("CATEGORY_DESCRIPTION"):
                    run.font.size = Pt(14)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                elif "_TITLE" in key:
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(68, 0, 136)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                elif "_DESCRIPTION" in key:
                    run.font.size = Pt(14)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            break


def get_matrix_markers(matrix_values: dict) -> dict:
    if not isinstance(matrix_values, dict):
        return {}

    matrix = matrix_values.get("matrix")
    if not isinstance(matrix, dict):
        matrix = matrix_values

    result = {}

    category_ids = range(1, 23)
    for cat_id in category_ids:
        category = ChildCategory.objects.filter(id=cat_id).first()
        if category:
            result[f"CATEGORY_TITLE_{category.id}"] = html.unescape(category.title)
            result[f"CATEGORY_DESCRIPTION_{category.id}"] = html.unescape(strip_tags(category.description))

    for marker, value in matrix.items():
        if marker not in MARKER_MODEL_MAP or value is None:
            continue

        key_base = marker.upper()
        for j, model_class in enumerate(MARKER_MODEL_MAP[marker]):
            obj = model_class.objects.filter(marker=marker, order_id=value).first()
            if obj:
                suffix = f"_{j}"
                title_with_number = f"{html.unescape(obj.title)} ({value})"
                result[f"{key_base}_TITLE{suffix}"] = title_with_number
                result[f"{key_base}_DESCRIPTION{suffix}"] = html.unescape(strip_tags(obj.description))

    matched_programs = matrix_values.get("matched_programs", [])
    for idx, program in enumerate(matched_programs[:10], start=1):
        result[f"PROGRAM_TITLE_{idx}"] = program.get("name", "")
        result[f"PROGRAM_DESC_{idx}"] = html.unescape(strip_tags(program.get("description", "")))

    for idx in range(1, 11):
        result.setdefault(f"PROGRAM_TITLE_{idx}", "")
        result.setdefault(f"PROGRAM_DESC_{idx}", "")

    return result


def fill_matrix_template(matrix_values, template_path, output_path):

    """
    Заполняет шаблон Word-файла данными из матрицы судьбы.
    """
    markers = get_matrix_markers(matrix_values)
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_text_with_style(paragraph, markers)

    for table in doc.tables:
        replace_text_in_table(table, markers)

    replace_text_in_headers(doc, markers)

    doc.save(output_path)
#