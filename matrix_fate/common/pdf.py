import os
import subprocess

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup


def generate_matrix_pdf(matrix_data, categories, base_image_path, output_dir):
    """
    Генерирует PDF-документ, содержащий:
    - изображение матрицы с нанесёнными значениями
    - текстовое описание категорий и талантов

    :param matrix_data: словарь значений матрицы (например, a, b, c...)
    :param categories: список категорий с описаниями и талантами
    :param base_image_path: путь до шаблона изображения (пустая матрица)
    :param output_dir: директория для сохранения выходных файлов
    :return: путь до итогового PDF
    """

    def render_matrix_on_image(base_image_path, matrix_data, output_image_path):
        image = Image.open(base_image_path).convert("RGB")
        draw = ImageDraw.Draw(image)
        font = ImageFont.truetype("DejaVuSans-Bold.ttf", size=20)

        # Пример координат — замените на реальные под вашу схему
        positions = {
            "a": (100, 100), "b": (150, 100), "c": (200, 100),
            "e": (125, 150), "g1": (250, 100), "p": (300, 100)
        }

        for key, (x, y) in positions.items():
            value = matrix_data.get(key)
            if value is not None:
                draw.text((x, y), str(value), font=font, fill="black")

        image.save(output_image_path)

    def strip_html(html):
        return BeautifulSoup(html, "html.parser").get_text(separator="\n")

    def generate_word_with_data(image_path, categories, output_path):
        doc = Document()
        doc.add_picture(image_path, width=Inches(5.5))
        doc.add_paragraph("\n")

        for category in categories:
            doc.add_heading(category['title'], level=1)
            doc.add_paragraph(strip_html(category.get('description', '')))

            for stage in ["birth_talent", "youth_talent", "mature_talent"]:
                talent = category.get(stage)
                if talent:
                    doc.add_paragraph(f"{talent['title']}: {talent['description']}")

            doc.add_paragraph("\n")

        doc.save(output_path)

    def convert_docx_to_pdf(docx_path):
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", docx_path
        ], check=True)

    os.makedirs(output_dir, exist_ok=True)
    filled_image_path = os.path.join(output_dir, "matrix_filled.png")
    docx_path = os.path.join(output_dir, "matrix_report.docx")

    render_matrix_on_image(base_image_path, matrix_data, filled_image_path)
    generate_word_with_data(filled_image_path, categories, docx_path)
    convert_docx_to_pdf(docx_path)

    return os.path.join(output_dir, "matrix_report.pdf")
